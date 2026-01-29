"""DOCX parser implementation."""

import hashlib
import re
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..models import Chunk, ChunkType, Heading, Locator
from .base import Parser


class DOCXParser(Parser):
    """Parser for DOCX files."""

    def parse(self, file_path: Path, doc_id: str, company_id: str) -> tuple[list[Chunk], list[Heading]]:
        """Parse a DOCX file."""
        doc = Document(str(file_path))
        chunks: list[Chunk] = []
        headings: list[Heading] = []
        
        structure_path: list[str] = []
        paragraph_index = 0

        for element in doc.element.body:
            if element.tag.endswith("p"):
                # Paragraph
                para = self._get_paragraph(doc, element)
                if para:
                    paragraph_index += 1
                    
                    # Check if it's a heading
                    if para.style.name.startswith("Heading"):
                        level = self._get_heading_level(para.style.name)
                        heading_text = para.text.strip()
                        
                        if heading_text:
                            # Update structure path
                            structure_path = structure_path[:level-1] + [heading_text]
                            
                            # Create heading
                            heading_id = hashlib.md5(
                                f"{doc_id}:{paragraph_index}".encode()
                            ).hexdigest()[:16]
                            
                            heading = Heading(
                                heading_id=heading_id,
                                doc_id=doc_id,
                                company_id=company_id,
                                heading_text=heading_text,
                                heading_norm=self._normalize_heading(heading_text),
                                level=level,
                                locator=Locator(paragraph_index=paragraph_index),
                            )
                            headings.append(heading)
                    
                    # Create chunk for paragraph
                    text = para.text.strip()
                    if text:
                        chunk_id = hashlib.md5(
                            f"{doc_id}:para:{paragraph_index}".encode()
                        ).hexdigest()[:16]
                        
                        chunk = Chunk(
                            chunk_id=chunk_id,
                            doc_id=doc_id,
                            company_id=company_id,
                            chunk_type=ChunkType.PARAGRAPH,
                            text=text,
                            structure_path=structure_path.copy(),
                            locator=Locator(paragraph_index=paragraph_index),
                        )
                        chunks.append(chunk)

            elif element.tag.endswith("tbl"):
                # Table
                table = self._get_table(doc, element)
                if table:
                    paragraph_index += 1
                    
                    # Extract table data
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    
                    # Create table text representation
                    table_text = "\n".join([" | ".join(row) for row in table_data])
                    
                    if table_text:
                        chunk_id = hashlib.md5(
                            f"{doc_id}:table:{paragraph_index}".encode()
                        ).hexdigest()[:16]
                        
                        chunk = Chunk(
                            chunk_id=chunk_id,
                            doc_id=doc_id,
                            company_id=company_id,
                            chunk_type=ChunkType.TABLE,
                            text=table_text,
                            structure_path=structure_path.copy(),
                            locator=Locator(paragraph_index=paragraph_index),
                            raw_table=table_data,
                        )
                        chunks.append(chunk)

        return chunks, headings

    def _get_paragraph(self, doc: Document, element) -> Paragraph | None:
        """Get paragraph from element."""
        for para in doc.paragraphs:
            if para._element == element:
                return para
        return None

    def _get_table(self, doc: Document, element) -> Table | None:
        """Get table from element."""
        for table in doc.tables:
            if table._element == element:
                return table
        return None

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        match = re.search(r"Heading (\d+)", style_name)
        if match:
            return int(match.group(1))
        return 1

    def _normalize_heading(self, text: str) -> str:
        """Normalize heading text for comparison."""
        # Remove numbers, special chars, convert to lowercase
        normalized = re.sub(r"[\d\.\-\s]+", " ", text.lower())
        normalized = re.sub(r"[^\w\s]", "", normalized)
        return " ".join(normalized.split())
