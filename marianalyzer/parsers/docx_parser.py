"""DOCX document parser using python-docx."""

import hashlib
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument

from marianalyzer.chunking.text_chunker import chunk_text
from marianalyzer.models import Chunk, Document, Heading, ParsedDocument
from marianalyzer.parsers.base import BaseParser
from marianalyzer.utils.citations import format_citation
from marianalyzer.utils.logging_config import get_logger

logger = get_logger()


class DOCXParser(BaseParser):
    """Parser for DOCX documents."""

    def supports_format(self, extension: str) -> bool:
        """Check if extension is .docx."""
        return extension.lower() == ".docx"

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse DOCX document.

        Args:
            file_path: Path to DOCX file

        Returns:
            ParsedDocument with chunks and headings
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        logger.info(f"Parsing DOCX: {file_path}")

        try:
            doc = DocxDocument(str(file_path))

            # Extract headings and paragraphs
            chunks = []
            headings = []
            chunk_index = 0
            paragraph_index = 0
            current_section = None

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()

                if not text:
                    continue

                # Check if paragraph is a heading
                if paragraph.style.name.startswith("Heading"):
                    level = self._get_heading_level(paragraph.style.name)

                    heading = Heading(
                        doc_id=0,  # Will be set when inserting to DB
                        level=level,
                        heading_text=text,
                        page_or_location=f"para_{paragraph_index}",
                    )
                    headings.append(heading)
                    current_section = text

                    # Also create a chunk for the heading
                    citation = format_citation(
                        file_path=str(file_path.name),
                        section=f"para_{paragraph_index}",
                    )

                    chunk = Chunk(
                        doc_id=0,
                        chunk_index=chunk_index,
                        chunk_text=text,
                        chunk_type="heading",
                        citation=citation,
                        metadata={
                            "paragraph_index": paragraph_index,
                            "heading_level": level,
                        },
                    )
                    chunks.append(chunk)
                    chunk_index += 1

                else:
                    # Regular paragraph - chunk it
                    text_chunks = chunk_text(text)

                    for chunk_text_content in text_chunks:
                        citation = format_citation(
                            file_path=str(file_path.name),
                            section=f"para_{paragraph_index}",
                        )

                        chunk = Chunk(
                            doc_id=0,
                            chunk_index=chunk_index,
                            chunk_text=chunk_text_content,
                            chunk_type="paragraph",
                            citation=citation,
                            metadata={
                                "paragraph_index": paragraph_index,
                                "section": current_section,
                            },
                        )
                        chunks.append(chunk)
                        chunk_index += 1

                paragraph_index += 1

            # Process tables
            for table_index, table in enumerate(doc.tables):
                table_chunks = self._parse_table(table, table_index, file_path.name)
                for table_chunk in table_chunks:
                    table_chunk.chunk_index = chunk_index
                    chunks.append(table_chunk)
                    chunk_index += 1

            logger.info(f"Extracted {len(chunks)} chunks and {len(headings)} headings")

            # Create document metadata
            file_hash = self._compute_file_hash(file_path)
            metadata = self._extract_metadata(doc)

            document = Document(
                file_path=str(file_path.name),
                file_hash=file_hash,
                file_type="docx",
                file_size=file_path.stat().st_size,
                metadata=metadata,
            )

            return ParsedDocument(
                metadata=document,
                chunks=chunks,
                headings=headings,
            )

        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        # Style name format: "Heading 1", "Heading 2", etc.
        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return 1

    def _parse_table(self, table: Any, table_index: int, file_name: str) -> list[Chunk]:
        """Parse a table into chunks (one per row)."""
        chunks = []

        # Get header row if it exists
        headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []

        # Process each row
        for row_index, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]

            # Skip empty rows
            if not any(row_data):
                continue

            # Create row text
            if row_index == 0:
                row_text = " | ".join(row_data)
            else:
                # Include headers for context
                row_text = " | ".join(f"{h}: {v}" for h, v in zip(headers, row_data))

            citation = format_citation(
                file_path=file_name,
                section=f"table_{table_index}_row_{row_index}",
            )

            chunk = Chunk(
                doc_id=0,
                chunk_index=0,  # Will be set by caller
                chunk_text=row_text,
                chunk_type="table_row",
                citation=citation,
                metadata={
                    "table_index": table_index,
                    "row_index": row_index,
                    "headers": headers,
                },
            )
            chunks.append(chunk)

        return chunks

    def _extract_metadata(self, doc: DocxDocument) -> dict[str, Any]:
        """Extract DOCX metadata."""
        metadata = {}

        if doc.core_properties:
            props = doc.core_properties
            metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
            }

        return metadata

    def _compute_file_hash(self, file_path: Path) -> str:
        """Compute SHA256 hash of file."""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
